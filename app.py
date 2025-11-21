import streamlit as st
import pandas as pd
from datetime import datetime
from fpdf import FPDF
import google.generativeai as genai
from PIL import Image
import io
import sqlite3
import hashlib
import json
import base64
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURATION & SETUP ---
st.set_page_config(
    page_title="SiteVision AI | Enterprise",
    page_icon="👁️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- BRANDING & CSS (UAT FIXES: Responsiveness, Branding Removal, Cleanliness) ---
def apply_custom_css():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
            color: #0F172A;
        }
        
        :root {
            --primary: #0F172A; /* Slate 900 */
            --accent: #3B82F6;  /* Electric Blue */
            --secondary: #64748B; /* Slate 500 */
            --bg: #F1F5F9;      /* Slate 100 */
            --card-bg: #FFFFFF;
        }

        .stApp {
            background-color: var(--bg);
        }
        
        /* 1. HIDE STREAMLIT BRANDING & FOOTER (UAT Fix) */
        footer {visibility: hidden;}
        /* Hide the hamburger menu */
        [data-testid="stSidebarContent"] > section > div.css-vk3y5n.e1fqn3694 { visibility: hidden; }
        /* Target the top-right menu */
        button[title="View source"] { display: none !important; }

        /* Sidebar Styling */
        [data-testid="stSidebar"] {
            background-color: var(--card-bg);
            border-right: 1px solid #E2E8F0;
        }
        
        /* 2. PROFESSIONAL CARD STYLING (For mobile and desktop) */
        div[data-testid*="stVerticalBlock"] > div[style*="flex-direction: column"] > div[data-testid*="stVerticalBlock"], 
        div.stContainer {
            background-color: var(--card-bg);
            padding: 24px;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -2px rgba(0, 0, 0, 0.05);
            border: 1px solid #E2E8F0;
            margin-bottom: 1.5rem;
        }
        
        /* Headers */
        h1, h2, h3 {
            color: var(--primary);
            font-weight: 700;
            letter-spacing: -0.02em;
        }
        
        /* Accent Headers */
        h3 {
            border-left: 4px solid var(--accent);
            padding-left: 12px;
            margin-top: 0.5rem;
            margin-bottom: 0.8rem;
        }
        
        /* Primary Buttons (Full width on mobile) */
        .stButton button {
            background-color: var(--accent);
            color: white;
            font-weight: 600;
            border-radius: 6px;
            border: none;
            height: 44px;
            transition: all 0.2s ease;
            width: 100%;
        }
        .stButton button:hover {
            background-color: #2563EB;
            box-shadow: 0 4px 12px rgba(59, 130, 246, 0.3);
            transform: translateY(-1px);
        }

        /* Metrics Styling */
        [data-testid="stMetricValue"] {
            color: var(--accent);
            font-weight: 700;
        }
        
        /* Form Inputs (Ensure full width) */
        .stTextInput, .stSelectbox, .stTextArea {
            width: 100% !important;
        }
        .stTextInput input, .stSelectbox div[data-baseweb="select"], .stTextArea textarea {
            border-radius: 6px;
            border: 1px solid #CBD5E1;
            padding: 10px;
        }
        
        /* Custom Defect Card for Draft Register */
        .defect-card {
            background: #F8FAFC;
            border-left: 4px solid #3B82F6; /* Use main accent for cleaner look */
            padding: 12px;
            margin-bottom: 8px;
            border-radius: 6px;
            font-size: 0.9em;
            box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        }
        
        /* Image hover for Inspection Page */
        .hover-zoom {
            transition: transform 0.3s ease;
            border-radius: 8px;
            cursor: zoom-in;
            border: 1px solid #E2E8F0;
            width: 100%; /* Ensure image fills container */
            max-width: 150px;
            height: auto;
            display: block;
        }
        .hover-zoom:hover {
            transform: scale(1.5);
            z-index: 1000;
            box-shadow: 0 10px 25px rgba(0,0,0,0.2);
        }
        
        /* Financial Summary Box */
        .financial-box {
            background-color: #EFF6FF;
            border: 1px solid #BFDBFE;
            border-radius: 8px;
            padding: 20px;
            text-align: center;
            margin-top: 10px;
        }
        .financial-total {
            font-size: 1.8em;
            font-weight: bold;
            color: #1E40AF;
        }
        
        /* Info/Warning Boxes - make them match the aesthetic */
        div[data-testid="stAlert"] {
            border-radius: 8px;
            padding: 12px;
        }
        
        /* Ensure columns stack nicely on small screens */
        @media (max-width: 768px) {
            div[data-testid*="stVerticalBlock"] > div[style*="flex-direction: column"] > div[data-testid*="stVerticalBlock"], 
            div.stContainer {
                padding: 16px;
                margin-bottom: 1rem;
            }
            .stSidebar {
                padding-top: 10px;
            }
        }
    </style>
    """, unsafe_allow_html=True)

def get_logo_svg():
    return """
    <svg width="100%" height="60" viewBox="0 0 250 60" fill="none" xmlns="http://www.w3.org/2000/svg">
        <rect x="10" y="10" width="40" height="40" rx="8" fill="#0F172A"/>
        <path d="M30 20L40 40H20L30 20Z" fill="#3B82F6"/>
        <circle cx="30" cy="32" r="3" fill="white"/>
        <text x="65" y="38" fill="#0F172A" font-family="Inter, sans-serif" font-weight="bold" font-size="24" letter-spacing="-0.5">SiteVision <tspan fill="#3B82F6">AI</tspan></text>
    </svg>
    """

# --- CONSTANTS & DATA ---
SEVERITY_LEVELS = [
    "Minor Defect (Maintenance)",
    "Major Defect (Structural/Significant)",
    "Safety Hazard (Immediate Action)",
    "Further Investigation Required"
]

AREAS = [
    "Site & Fencing", "Exterior Walls", "Sub-floor Space", "Roof Exterior", 
    "Roof Space", "Interior", "Garage/Carport", "Wet Areas", "Outbuildings"
]

# --- HELPER FUNCTIONS ---
def parse_cost(cost_str):
    """Extracts low and high values from strings like '$500 - $1,200'"""
    if not cost_str or cost_str == "N/A":
        return 0, 0
    # Find all numbers (removing commas and non-digits)
    nums = re.findall(r'\d+', cost_str.replace(',', ''))
    if not nums:
        return 0, 0
    nums = [int(n) for n in nums]
    if len(nums) == 1:
        return nums[0], nums[0]
    return min(nums), max(nums)

def calculate_total_repairs(defects):
    total_min = 0
    total_max = 0
    for d in defects:
        c_min, c_max = parse_cost(d.get('cost', 'N/A'))
        total_min += c_min
        total_max += c_max
    return total_min, total_max

# --- DATABASE (SQLite for Auth & Reports) ---

def get_db_connection():
    return sqlite3.connect('sitevision.db')

def init_db():
    conn = get_db_connection()
    c = conn.cursor()
    # 1. Users Table for Auth
    c.execute('''CREATE TABLE IF NOT EXISTS users (username TEXT PRIMARY KEY, password TEXT, role TEXT, full_name TEXT)''')
    # 2. Reports Table for Save/Load/Version Control
    c.execute('''
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            report_title TEXT,
            address TEXT,
            inspector TEXT,
            saved_at TEXT,
            report_data_json TEXT
        )
    ''')
    
    # Default admin user
    secure_pass = hashlib.sha256(b"inspect").hexdigest()
    c.execute("SELECT * FROM users WHERE username = 'admin'")
    if not c.fetchone():
        c.execute("INSERT INTO users VALUES ('admin', ?, 'admin', 'System Administrator')", (secure_pass,))
    else:
        # Ensures the default password is set if needed for testing
        c.execute("UPDATE users SET password = ? WHERE username = 'admin'", (secure_pass,))
        
    conn.commit()
    conn.close()

# Auth Functions
def check_login(username, password):
    conn = get_db_connection()
    c = conn.cursor()
    hashed = hashlib.sha256(password.encode()).hexdigest()
    c.execute("SELECT role, full_name FROM users WHERE username = ? AND password = ?", (username, hashed))
    user = c.fetchone()
    conn.close()
    return user

def get_all_users():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT username, role, full_name FROM users")
    users = c.fetchall()
    conn.close()
    return users

def add_new_user(username, password, role, full_name):
    conn = get_db_connection()
    c = conn.cursor()
    try:
        hashed = hashlib.sha256(password.encode()).hexdigest()
        c.execute("INSERT INTO users VALUES (?, ?, ?, ?)", (username, hashed, role, full_name))
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        conn.close()
        return False

def delete_user(username):
    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute("DELETE FROM users WHERE username = ?", (username,))
        conn.commit()
        conn.close()
        return True
    except Exception:
        conn.close()
        return False

# Report Management Functions
def save_report(title, address, inspector, data):
    conn = get_db_connection()
    c = conn.cursor()
    saved_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    data_json = json.dumps(data)
    
    c.execute("INSERT INTO reports (report_title, address, inspector, saved_at, report_data_json) VALUES (?, ?, ?, ?, ?)", 
              (title, address, inspector, saved_at, data_json))
    conn.commit()
    conn.close()
    return saved_at

def get_all_reports(inspector=None):
    conn = get_db_connection()
    c = conn.cursor()
    if inspector and st.session_state.get('role') != 'admin':
        c.execute("SELECT id, report_title, address, inspector, saved_at FROM reports WHERE inspector = ? ORDER BY saved_at DESC", (inspector,))
    else:
        c.execute("SELECT id, report_title, address, inspector, saved_at FROM reports ORDER BY saved_at DESC")
    reports = c.fetchall()
    conn.close()
    return reports

def load_report_data(report_id):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT report_data_json FROM reports WHERE id = ?", (report_id,))
    data = c.fetchone()
    conn.close()
    if data:
        return json.loads(data[0])
    return None

def delete_report(report_id):
    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute("DELETE FROM reports WHERE id = ?", (report_id,))
        conn.commit()
        conn.close()
        return True
    except Exception:
        conn.close()
        return False

# --- AI ENGINE (UAT FIX: Improved Australian Property Lookup Prompt) ---
class AIEngine:
    def __init__(self, api_key):
        self.api_key = api_key
        self.client = None
        self.model_name = 'gemini-2.5-flash'
        if api_key:
            genai.configure(api_key=api_key)
            self.client = genai.Client()

    def _generate_content_text(self, prompt, system_prompt=None):
        if not self.client: return "AI Error: API Key not configured."
        
        config = {}
        if system_prompt:
            config['system_instruction'] = system_prompt
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt,
                config=config
            )
            return response.text
        except Exception as e:
            st.error(f"AI Generation Failed: {e}")
            return f"AI Generation Failed: {e}"

    def analyze_photo(self, image):
        if not self.client: return None
        prompt = """
        Act as an Australian Building Inspector. Analyze this image for defects and compliance against AS 4349.1.
        Format ONLY the output as follows, using Australian terminology:
        Defect: [Name, e.g., Spalling Concrete]
        Observation: [Technical Description of the condition]
        Recommendation: [Required action, e.g., Engage structural engineer for repair]
        """
        try: return self.client.models.generate_content(model=self.model_name, contents=[prompt, image]).text
        except: return "AI Error: Failed to analyze image."

    def get_property_history(self, address):
        if not self.client: return None, "N/A"
        
        # UAT FIX: More specific prompt for reliable Australian data extraction
        system_prompt = """
        You are a specialized data extractor for Australian property records. Given the search results for an Australian property address, find and return ONLY the following two pieces of information, formatted as a JSON string: 
        {'year_built': [Four digit year (e.g., '1985') or 'N/A'], 'property_type': [The most specific type, e.g., 'House', 'Unit', 'Townhouse', 'Apartment', or 'N/A']}.
        Base your answer ONLY on the provided search results from reliable Australian real estate or government websites (like Domain, Realestate.com.au). Do not add any text before or after the JSON.
        """

        prompt = f"Using Google Search, find the official year built and the property classification for the Australian property located at: {address}."
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt,
                config=genai.types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    tools=[{"google_search": {} }],
                    response_mime_type="application/json",
                    response_schema={"type": "OBJECT", "properties": {"year_built": {"type": "STRING"}, "property_type": {"type": "STRING"}}}
                )
            )
            json_str = response.text.strip()
            data = json.loads(json_str)
            
            year_str = data.get('year_built', 'N/A')
            p_type = data.get('property_type', 'N/A')
            
            try: year = int(year_str) if year_str.isdigit() and len(year_str) == 4 else None
            except: year = None
                
            return year, p_type
            
        except Exception as e:
            print(f"Error fetching property history: {e}")
            return None, "AI Search Failed"

    def generate_liability_statement(self, defect, severity):
        prompt = f"Act as an Australian legal risk assessor for building inspections. Draft a concise statement (2-3 sentences max) outlining the potential legal liability, litigation risk, and insurance implications for the property owner if the '{defect}' (Severity: {severity}) is not rectified promptly. Reference duty of care and Australian consumer law briefly."
        return self._generate_content_text(prompt)

    def check_compliance(self, query):
        if not self.client: return "AI Error: API Key not configured."
        system_prompt = "You are a compliance officer. Search and cite the specific clause from the NCC 2022 Vol 2 or relevant Australian Standards (AS 4349.1, AS 3740, etc.) that governs the following query. If a specific clause is not found, state the general principle and its potential impact."
        
        prompt = f"Compliance query: {query}"
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt,
                config=genai.types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    tools=[{"google_search": {} }]
                )
            )
            return response.text
        except: return "AI Compliance Check Failed"

    def estimate_cost(self, defect, severity):
        prompt = f"Provide a repair cost range in AUD for '{defect}' ({severity}) in Australia. Return ONLY the range string (e.g. '$500 - $1,000'). Do not add any other text."
        return self._generate_content_text(prompt)
    
    def generate_scope(self, defect, rec):
        prompt = f"Write a detailed Scope of Works for a tradesperson to rectify '{defect}'. Recommendation was: '{rec}'. Use Australian trade terminology."
        return self._generate_content_text(prompt)

    def explain_impact(self, defect):
        prompt = f"Explain the consequences to the property owner if '{defect}' is not fixed. Focus on structural integrity, cost escalation, and compliance risk under Australian law."
        return self._generate_content_text(prompt)

    def suggest_trade(self, defect):
        prompt = f"Which specific Australian licensed trade is best suited to fix '{defect}'? (e.g. Roof Plumber, Electrician, Structural Engineer). Return ONLY the trade name."
        return self._generate_content_text(prompt)

    def generate_maintenance(self, age, defects):
        d_list = ", ".join([d['defect_name'] for d in defects])
        prompt = f"Create a 5-year maintenance schedule for a house built in {age} with these current defects: {d_list}. Format as a markdown list."
        return self._generate_content_text(prompt)
    
    def generate_exec_summary(self, defects, total_cost):
        d_list = ", ".join([f"{d['defect_name']} ({d['severity']})" for d in defects])
        prompt = f"Write an Executive Summary for a building report compliant with AS 4349.1. Defects: {d_list}. Total Estimated Rectification Cost: {total_cost}. Focus on major risks and required immediate action."
        return self._generate_content_text(prompt)


# --- EXPORT ENGINES (PDF/DOCX) ---
# (Export functions remain largely the same, ensuring compatibility with the data structure)

class ReportPDF(FPDF):
    def __init__(self, company, license, logo_path, header_img, footer_img):
        super().__init__()
        self.company = company
        self.license = license
        self.logo_path = logo_path
        self.header_img = header_img
        self.footer_img = footer_img

    def header(self):
        if self.header_img:
            try: self.image(self.header_img, 0, 0, 210)
            except: pass
            
        if self.logo_path and not self.header_img:
            try: self.image(self.logo_path, 10, 8, 35)
            except: pass
            
        self.ln(25)
        if not self.header_img:
            self.set_font('Arial', 'B', 16)
            self.cell(0, 10, f"{self.company} - Inspection Report", 0, 1, 'R')
            self.line(10, 35, 200, 35)
        self.ln(10)

    def footer(self):
        if self.footer_img:
            try: self.image(self.footer_img, 0, 270, 210)
            except: pass
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, 'C')

def generate_pdf(data, prop, inspector, co_details, summary, total_est_str):
    logo = io.BytesIO(co_details['logo'].getvalue()) if co_details['logo'] else None
    hed = io.BytesIO(co_details['header'].getvalue()) if co_details['header'] else None
    foot = io.BytesIO(co_details['footer'].getvalue()) if co_details['footer'] else None

    pdf = ReportPDF(co_details['name'], co_details['lic'], logo, hed, foot)
    pdf.add_page()
    
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, "Property Overview", 0, 1)
    pdf.set_font('Arial', '', 11)
    pdf.cell(40, 7, "Address:", 0); pdf.cell(0, 7, prop['address'], 0, 1)
    pdf.cell(40, 7, "Client:", 0); pdf.cell(0, 7, prop['client'], 0, 1)
    pdf.cell(40, 7, "Year Built:", 0); pdf.cell(0, 7, str(prop['year_built']), 0, 1)
    pdf.cell(40, 7, "Property Type:", 0); pdf.cell(0, 7, prop['property_type'], 0, 1)
    pdf.ln(5)
    
    # Executive Summary
    pdf.set_fill_color(240, 245, 255)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, "  Executive Summary", 0, 1, 'L', 1)
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 6, summary)
    pdf.ln(5)

    # Financial Summary (The Calculator Result)
    pdf.set_fill_color(255, 247, 237) # Light Orange
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, "  Financial Estimate", 0, 1, 'L', 1)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 10, f"Total Estimated Rectification Costs: {total_est_str}", 0, 1)
    pdf.set_font('Arial', 'I', 9)
    pdf.multi_cell(0, 5, "Note: These costs are rough estimates generated by AI based on average Australian trade rates and edited by the inspector. Actual quotes from licensed trades should be sought.")
    pdf.ln(8)
    
    # Detailed Defects
    for item in data:
        pdf.set_font('Arial', 'B', 11)
        if "Safety" in item['severity']: pdf.set_text_color(200, 0, 0)
        elif "Major" in item['severity']: pdf.set_text_color(200, 100, 0)
        else: pdf.set_text_color(0, 50, 150)
        
        pdf.cell(0, 8, f"{item['area']} | {item['defect_name']}", 0, 1)
        pdf.set_text_color(0)
        
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, f"Observation: {item['observation']}")
        pdf.multi_cell(0, 5, f"Rectification: {item['recommendation']}")
        
        # Cost Line
        pdf.set_font('Arial', 'B', 9)
        pdf.cell(30, 6, "Est. Repair Cost:", 0)
        pdf.set_font('Arial', '', 9)
        pdf.cell(0, 6, f"{item.get('cost', 'N/A')}", 0, 1)
        
        if 'scope' in item and item['scope']:
            pdf.set_font('Arial', 'I', 9)
            pdf.multi_cell(0, 5, f"Scope of Works: {item['scope']}")
            
        if 'impact' in item and item['impact']:
            pdf.set_font('Arial', 'I', 9)
            pdf.multi_cell(0, 5, f"Impact Analysis: {item['impact']}")

        if 'liability' in item and item['liability']:
            pdf.set_font('Arial', 'BI', 9) # Bold and Italic for legal warning
            pdf.set_text_color(150, 0, 0) # Dark Red for liability
            pdf.multi_cell(0, 5, f"Legal Liability / Risk Statement: {item['liability']}") 
            pdf.set_text_color(0) # Reset color
            
        pdf.ln(4)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
        
    return pdf.output(dest='S').encode('latin-1')

def generate_docx(data, prop, inspector, co_details, summary, total_est_str):
    doc = Document()
    
    # Header
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.autofit = False
    htable.columns[0].width = Inches(2)
    htable.columns[1].width = Inches(4)
    
    if co_details['logo']:
        logo_stream = io.BytesIO(co_details['logo'].getvalue())
        htable.cell(0,0).paragraphs[0].add_run().add_picture(logo_stream, width=Inches(1.5))
    
    htable.cell(0,1).text = f"{co_details['name']}\nLicence: {co_details['lic']}\nAS 4349.1 Inspection Report"
    
    # Body
    doc.add_heading('Property Inspection Report', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Address: {prop['address']}\n").bold = True
    p.add_run(f"Client: {prop['client']}\n")
    p.add_run(f"Year Built: {prop['year_built']}\n")
    p.add_run(f"Property Type: {prop['property_type']}\n")
    p.add_run(f"Inspector: {inspector}\n")
    p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(summary)

    # Financial Section (The Calculator Result)
    doc.add_heading('Estimated Repair Costs', level=1)
    p_fin = doc.add_paragraph()
    p_fin.add_run(f"Total Estimate: {total_est_str}").bold = True
    p_fin.add_run("\n*These estimates are indicative only and subject to trade quotes.").italic = True
    
    doc.add_heading('Defect Register', level=1)
    
    for item in data:
        # Using a table to align image and text side-by-side (responsive layout for docx)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        img_cell = table.cell(0, 0)
        paragraph = img_cell.paragraphs[0]
        if 'image_data' in item and item['image_data']:
            try:
                img_bytes = base64.b64decode(item['image_data'])
                img_stream = io.BytesIO(img_bytes)
                paragraph.add_run().add_picture(img_stream, width=Inches(2.5)) 
            except:
                paragraph.text = "[Image Error]"
        else:
            paragraph.text = "[No Image]"
            
        txt_cell = table.cell(0, 1)
        txt_cell.text = ""
        p = txt_cell.add_paragraph()
        p.add_run(f"{item['defect_name']}").bold = True
        p.add_run(f"\nArea: {item['area']}")
        severity_run = p.add_run(f"\nSeverity: {item['severity']}")
        if "Major" in item['severity'] or "Safety" in item['severity']:
             severity_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        
        p.add_run("\n\nObservation:").bold = True
        p.add_run(f" {item['observation']}")
        
        p.add_run("\n\nRecommendation:").bold = True
        p.add_run(f" {item['recommendation']}")
        
        p.add_run("\n\nEst. Cost: ").bold = True
        p.add_run(f"{item.get('cost', 'N/A')}")

        if item.get('scope'):
            p.add_run("\nScope of Works: ").bold = True
            p.add_run(f"{item['scope']}")
        
        if item.get('impact'):
            p.add_run("\nImpact Analysis: ").bold = True
            p.add_run(f"{item['impact']}")

        if item.get('liability'):
            p.add_run("\n\nLegal Liability / Risk Statement: ").bold = True
            p.add_run(f"{item['liability']}").italic = True

        doc.add_paragraph("")

    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

# --- UI MODULES ---
def login_page():
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        st.markdown(get_logo_svg(), unsafe_allow_html=True)
        st.markdown("""
        <div style='background:white; padding:40px; border-radius:12px; box-shadow:0 4px 20px rgba(0,0,0,0.05); margin-top:20px; text-align:center; border: 1px solid #E2E8F0;'>
            <h3 style='color:#0F172A; border-left: none; padding-left: 0;'>Inspector Portal</h3>
            <p style='color:#64748B; font-size:0.9em; margin-bottom: 20px;'>Enterprise Edition v6.1 | Use 'admin' / 'inspect'</p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("login_form"):
            u = st.text_input("Username", help="Enter your registered ID")
            p = st.text_input("Password", type="password")
            st.markdown("<br>", unsafe_allow_html=True)
            if st.form_submit_button("Secure Login"):
                user = check_login(u, p)
                if user:
                    st.session_state.update({
                        'logged_in': True, 
                        'role': user[0], 
                        'fullname': user[1], 
                        'page': 'New Inspection'
                    })
                    if 'year_built' not in st.session_state: st.session_state['year_built'] = 2000
                    if 'property_type' not in st.session_state: st.session_state['property_type'] = 'N/A'
                    st.rerun()
                else: st.error("Access Denied")

def inspection_page(ai: AIEngine):
    section_header("New Inspection: Data Capture", "clipboard")
    
    # --- Property Details and History Lookup ---
    with st.container():
        st.subheader("📍 Site Inspection Details") # Australian Terminology
        
        c1, c2 = st.columns(2)
        st.session_state['addr'] = c1.text_input("Address", st.session_state.get('addr', ''))
        st.session_state['client'] = c2.text_input("Client", st.session_state.get('client', ''))
        
        st.session_state['year_built'] = st.number_input("Year Built (Manual/AI)", 1900, 2025, st.session_state.get('year_built', 2000))
        st.info(f"Property Type (AI Lookup): **{st.session_state.get('property_type', 'N/A')}**")
        
        
        if st.button("Lookup Property History (AI)", use_container_width=True):
            if st.session_state['addr']:
                with st.spinner(f"Searching Australian property records for {st.session_state['addr']}..."):
                    fetched_year, fetched_type = ai.get_property_history(st.session_state['addr'])
                    if fetched_year or fetched_type != 'N/A':
                        if fetched_year: st.session_state['year_built'] = fetched_year
                        if fetched_type != 'N/A': st.session_state['property_type'] = fetched_type
                        st.success(f"History Found! Year Built: {st.session_state['year_built']}, Type: {st.session_state['property_type']}")
                    else:
                        st.warning("Could not retrieve reliable details. Please verify the address or enter manually.")
            else:
                st.warning("Please enter an Address first.")
        
        if st.session_state['year_built'] and st.button("Generate 5-Year Maintenance Plan (AI)", key="gen_maint_plan", use_container_width=True):
            with st.spinner("Building 5-Year Plan..."):
                plan = ai.generate_maintenance(st.session_state['year_built'], st.session_state['defects'])
                st.session_state['maint_plan'] = plan
                st.success("Maintenance Plan Generated (View on Dashboard)")

    st.markdown("<br>", unsafe_allow_html=True)
    
    # --- Main Inspection Columns (Stackable on Mobile) ---
    col_main, col_sidebar = st.columns([2, 1])
    
    with col_main:
        section_header("Defect & Compliance Entry", "camera")
        
        # --- AI Compliance Checker Section ---
        with st.expander("🔍 AI Compliance Checker (NCC/AS Lookup)", expanded=False):
            query = st.text_input("Enter a query for compliance search (e.g., 'Minimum balcony balustrade height in NCC 2022')")
            if st.button("Search Standards", key="compliance_search"):
                if query:
                    with st.spinner("Searching and citing standards..."):
                        result = ai.check_compliance(query)
                        st.session_state['compliance_result'] = result
                else:
                    st.warning("Please enter a query.")
                    
            if st.session_state.get('compliance_result'):
                st.markdown("#### Compliance Report")
                st.info(st.session_state['compliance_result'])

        st.markdown("<hr>", unsafe_allow_html=True)

        with st.container():
            area = st.selectbox("Area Inspected", AREAS)
            
            st.markdown("##### 📸 Evidence")
            c_img, c_analyze = st.columns([1, 2])
            img_file = c_img.file_uploader("Upload Photo", type=['jpg', 'png'], label_visibility="collapsed")
            
            b64_str = None
            if img_file:
                b = img_file.getvalue()
                b64_str = base64.b64encode(b).decode()
                c_img.markdown(f'<img src="data:image/png;base64,{b64_str}" class="hover-zoom">', unsafe_allow_html=True)
                
                if c_analyze.button("Analyze & Pre-fill Details (AI)", key="analyze_defect", use_container_width=True):
                    with st.spinner("Checking AS 4349.1..."):
                        ai_data = ai.analyze_photo(Image.open(img_file))
                        st.session_state['temp_ai'] = ai_data

            # Pre-fill fields if AI data exists from analysis
            d_d, d_o, d_r, sev_default = "", "", "", 0
            if 'temp_ai' in st.session_state and st.session_state['temp_ai']:
                raw = st.session_state['temp_ai']
                if "Defect:" in raw: d_d = raw.split("Defect:", 1)[1].split("\n")[0].strip()
                if "Observation:" in raw: d_o = raw.split("Observation:", 1)[1].split("\n")[0].strip()
                if "Recommendation:" in raw: d_r = raw.split("Recommendation:", 1)[1].split("\n")[0].strip()
                if "Safety" in raw: sev_default = 2
                elif "Major" in raw: sev_default = 1
                
            with st.form("defect"):
                name = st.text_input("Defect Title", value=d_d)
                
                c_ob, c_re = st.columns(2)
                obs = c_ob.text_area("Observation", value=d_o, height=100)
                rec = c_re.text_area("Recommendation", value=d_r, height=100)
                
                sev = st.selectbox("Severity", SEVERITY_LEVELS, index=sev_default)
                
                st.markdown("##### AI Context Generation")
                c_ai_1, c_ai_2, c_ai_3, c_ai_4 = st.columns(4)
                want_scope = c_ai_1.checkbox("Scope of Works")
                want_impact = c_ai_2.checkbox("Impact Analysis")
                want_trade = c_ai_3.checkbox("Suggest Trade")
                want_liability = c_ai_4.checkbox("Legal Liability Statement") # UAT requested feature

                if st.form_submit_button("Save Defect & Run AI Analysis"):
                    scope_txt, impact_txt, trade_txt, liability_txt = "", "", "", ""
                    
                    if not name or not obs or not rec:
                        st.error("Please fill in Defect Title, Observation, and Recommendation.")
                        st.stop()
                        
                    # Estimate cost automatically (MANDATORY)
                    cost_est_val = "N/A"
                    with st.spinner("1. Estimating Cost (AI)..."):
                        try: cost_est_val = ai.estimate_cost(name, sev)
                        except Exception: pass
                            
                    # Run other requested AI tasks
                    with st.spinner("2. Generating supplementary AI analysis..."):
                        if want_scope: scope_txt = ai.generate_scope(name, rec)
                        if want_impact: impact_txt = ai.explain_impact(name)
                        if want_trade: trade_txt = ai.suggest_trade(name)
                        if want_liability: liability_txt = ai.generate_liability_statement(name, sev)
                    
                    st.session_state['defects'].append({
                        "area": area, "defect_name": name, "observation": obs,
                        "severity": sev, "recommendation": rec,
                        "scope": scope_txt, "impact": impact_txt, "trade": trade_txt,
                        "liability": liability_txt,
                        "cost": cost_est_val,
                        "image_data": b64_str
                    })
                    st.success(f"Defect '{name}' saved to Draft (Cost: {cost_est_val})")
                    st.session_state['temp_ai'] = None
                    st.rerun()

    with col_sidebar:
        st.subheader("📋 Draft Register")
        if st.session_state['defects']:
            t_min, t_max = calculate_total_repairs(st.session_state['defects'])
            total_str = f"${t_min:,} - ${t_max:,}"
            st.markdown(f"**Total Draft Est:** <span style='color:#3B82F6;'>{total_str}</span>", unsafe_allow_html=True)
            st.markdown("<hr style='margin: 8px 0;'>", unsafe_allow_html=True)
            
            for i, d in enumerate(st.session_state['defects']):
                st.markdown(f"""
                <div class="defect-card">
                    <b>{d['area']}</b><br>
                    {d['defect_name']}<br>
                    <span style="color:#EF4444; font-size:0.8em;">{d['severity'].split('(')[0]}</span><br>
                    <span style="color:#059669; font-size:0.8em;">Est: {d.get('cost','N/A')}</span>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No defects logged.")
            
def manage_reports_page():
    section_header("Load & Manage Reports (Version Control)", "archive")
    
    st.info("View, load, or delete saved report versions. Only admins can see all reports; inspectors see their own.")
    
    reports = get_all_reports(inspector=st.session_state.get('fullname'))
    
    if not reports:
        st.warning("No saved reports found in the database for this user.")
        return

    df_reports = pd.DataFrame(reports, columns=['ID', 'Title', 'Address', 'Inspector', 'Saved At'])
    
    st.subheader("Saved Report Versions")
    st.dataframe(df_reports, use_container_width=True, hide_index=True)

    col1, col2 = st.columns(2)
    
    # Load Functionality
    report_ids = df_reports['ID'].tolist()
    load_id = col1.selectbox("Select Report ID to Load", report_ids)
    
    if col1.button("Load Report", use_container_width=True):
        if load_id:
            data = load_report_data(load_id)
            if data:
                # Load data back into session state
                for key, value in data.items():
                    st.session_state[key] = value
                
                st.session_state['page'] = 'Finalize Report'
                st.success(f"Report ID {load_id} loaded successfully. Redirecting to Finalize Report.")
                st.rerun()
            else:
                st.error("Error loading report data.")

    # Delete Functionality
    delete_id = col2.selectbox("Select Report ID to Delete", report_ids)
    if col2.button("Delete Report", use_container_width=True):
        if delete_id:
            if delete_report(delete_id):
                st.success(f"Report ID {delete_id} deleted permanently.")
                st.rerun()
            else:
                st.error("Failed to delete report.")

def report_page(ai: AIEngine):
    section_header("Finalize Report: Review & Export", "file-text")
    
    if not st.session_state['defects']:
        st.warning("No defects logged. Please start a new inspection or load a draft from 'Manage Reports'.")
        return
        
    # Branding Options
    with st.expander("🎨 Report Branding & Maintenance Plan", expanded=False):
        c1, c2, c3 = st.columns(3)
        st.session_state['header_img'] = c1.file_uploader("Header Image (PDF/DOCX)", type=['png', 'jpg'])
        st.session_state['footer_img'] = c2.file_uploader("Footer Image (PDF)", type=['png', 'jpg'])
        if st.session_state.get('maint_plan'):
            with c3.expander("View 5-Year Maintenance Plan"):
                st.markdown(st.session_state['maint_plan'])
        else:
            c3.info("No Maintenance Plan Generated.")

    st.subheader("Final Defect Register Review & Costing")
    st.info("You can make final manual edits (e.g., adjust cost or wording) directly in the table below.")
    
    # Convert defects list to DataFrame for the data editor
    df = pd.DataFrame(st.session_state['defects'])
    
    # Use data editor to allow inspector to make final manual edits (especially to cost)
    edited_df = st.data_editor(
        df, 
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "severity": st.column_config.SelectboxColumn("Severity", options=SEVERITY_LEVELS),
            "cost": st.column_config.TextColumn("Est Cost ($)", help="Enter range e.g. $500 - $1,000")
        },
        hide_index=True,
        # Ensure all columns are included for completeness
        column_order=["area", "defect_name", "severity", "observation", "recommendation", "trade", "cost", "scope", "impact", "liability", "image_data"]
    )
    st.session_state['defects'] = edited_df.to_dict('records')
    
    # CALCULATE TOTALS from the edited data
    t_min, t_max = calculate_total_repairs(st.session_state['defects'])
    total_str = f"${t_min:,} - ${t_max:,}"
    
    st.markdown(f"""
    <div class="financial-box">
        <h3>💰 Total Estimated Rectification Costs</h3>
        <div class="financial-total">{total_str}</div>
        <p><i>Based on the current register. Costs are indicative only.</i></p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<hr>", unsafe_allow_html=True)

    st.subheader("Executive Summary")
    
    c_sum, c_gen = st.columns([3, 1])
    if c_gen.button("Generate Summary (AI)", use_container_width=True):
        st.session_state['summary'] = ai.generate_exec_summary(st.session_state['defects'], total_str)
        
    summ = c_sum.text_area("Executive Summary (Final Edit)", st.session_state.get('summary', ''), height=200, label_visibility="collapsed")

    st.markdown("### Report Management & Export")
    c_save, c_pdf, c_docx = st.columns(3)
    
    # --- SAVE / VERSION CONTROL ---
    report_title = c_save.text_input("Report Title for Saving", value=f"Report for {st.session_state.get('addr', 'New Property')}")
    if c_save.button("Save Current Draft (New Version)", use_container_width=True):
        if report_title and st.session_state.get('addr'):
            # Data to save (only critical data needed for a report)
            data_to_save = {
                'addr': st.session_state.get('addr'),
                'client': st.session_state.get('client'),
                'year_built': st.session_state.get('year_built'),
                'property_type': st.session_state.get('property_type'),
                'defects': st.session_state['defects'],
                'summary': summ,
                'maint_plan': st.session_state.get('maint_plan', '')
            }
            saved_time = save_report(report_title, st.session_state['addr'], st.session_state['fullname'], data_to_save)
            st.success(f"Report '{report_title}' saved successfully! Version: {saved_time}")
        else:
            st.error("Please enter an Address and Report Title before saving.")
        
    # --- EXPORT PREPARATION ---
    property_data = {
        "address": st.session_state.get('addr', ''), 
        "client": st.session_state.get('client', ''),
        "year_built": st.session_state.get('year_built', 'N/A'),
        "property_type": st.session_state.get('property_type', 'N/A'),
    }

    company_data = {
        "name": st.session_state.get('co_name', ''), 
        "lic": st.session_state.get('lic', ''), 
        "logo": st.session_state.get('logo_file'), 
        "header": st.session_state.get('header_img'), 
        "footer": st.session_state.get('footer_img')
    }

    # PDF Export
    pdf_dat = generate_pdf(
        st.session_state['defects'],
        property_data,
        st.session_state['fullname'],
        company_data,
        summ,
        total_str
    )
    c_pdf.download_button(
        "Download PDF Report", 
        pdf_dat, 
        f"Inspection_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf", 
        "application/pdf",
        use_container_width=True
    )
        
    # DOCX Export
    docx_dat = generate_docx(
        st.session_state['defects'],
        property_data,
        st.session_state['fullname'],
        company_data,
        summ,
        total_str
    )
    c_docx.download_button(
        "Download Word (DOCX)", 
        docx_dat, 
        f"Inspection_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

def admin_page():
    section_header("Administration Panel", "lock")
    
    if st.session_state['role'] != 'admin':
        st.error("Access Restricted to System Administrators.")
        return

    st.subheader("Manage User Accounts")
    
    with st.expander("➕ Add New User", expanded=False):
        with st.form("add_user"):
            c1, c2 = st.columns(2)
            u = c1.text_input("Username (Login ID)")
            fn = c2.text_input("Full Name")
            p = c1.text_input("Password", type="password")
            r = c2.selectbox("Role", ["inspector", "admin"])
            if st.form_submit_button("Create User", use_container_width=True):
                if u and p and fn:
                    if add_new_user(u, p, r, fn):
                        st.success(f"User '{u}' created successfully with role '{r}'.")
                        st.rerun()
                    else:
                        st.error("Username already exists or database error.")
                else:
                    st.warning("All fields are required.")

    st.subheader("Existing Users")
    users = get_all_users()
    if users:
        df_users = pd.DataFrame(users, columns=['Username', 'Role', 'Full Name'])
        st.dataframe(df_users, use_container_width=True, hide_index=True)
        
        st.markdown("---")
        with st.form("delete_user"):
            del_options = [user[0] for user in users if user[0] != 'admin']
            user_to_delete = st.selectbox("Select User to Delete", del_options)
            
            if st.form_submit_button("Delete Selected User", use_container_width=True):
                if user_to_delete:
                    if delete_user(user_to_delete):
                        st.success(f"User '{user_to_delete}' deleted.")
                        st.rerun()
                    else:
                        st.error("Failed to delete user.")
                else:
                    st.warning("Please select a user to delete.")
    else:
        st.info("No users registered yet.")

def section_header(text, icon):
    st.markdown(f"""
    <div style="margin-bottom: 20px; border-bottom: 2px solid #E2E8F0; padding-bottom: 10px;">
        <h2 style="margin: 0; color: #0F172A;">{text}</h2>
    </div>
    """, unsafe_allow_html=True)

# --- MAIN APPLICATION ENTRY POINT ---
def main():
    init_db()
    apply_custom_css()
    
    # Initialize session state variables
    if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False
    if 'defects' not in st.session_state: st.session_state['defects'] = []
    if 'page' not in st.session_state: st.session_state['page'] = 'New Inspection'
    if 'year_built' not in st.session_state: st.session_state['year_built'] = 2000
    if 'property_type' not in st.session_state: st.session_state['property_type'] = 'N/A'
    if 'summary' not in st.session_state: st.session_state['summary'] = ''
    
    # Initialize AI Engine (needs API key from settings)
    ai = AIEngine(st.session_state.get('api_key'))
    
    if not st.session_state['logged_in']:
        login_page()
    else:
        # Sidebar Navigation
        with st.sidebar:
            st.markdown(get_logo_svg(), unsafe_allow_html=True)
            st.markdown(f"**Inspector:** {st.session_state.get('fullname')}", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            
            menu = ["New Inspection", "Manage Reports", "Dashboard", "Finalize Report"]
            if st.session_state['role'] == 'admin':
                menu.append("Admin")
                
            pg = st.radio("Navigate", menu, index=menu.index(st.session_state['page']) if st.session_state['page'] in menu else 0)
            st.session_state['page'] = pg
            
            with st.expander("⚙️ Report Settings", expanded=False):
                st.session_state['api_key'] = st.text_input("Gemini API Key", type="password", value=st.session_state.get('api_key', ''))
                st.session_state['co_name'] = st.text_input("Company Name", value=st.session_state.get('co_name', 'SiteVision Pty Ltd'))
                st.session_state['lic'] = st.text_input("Inspector Licence", value=st.session_state.get('lic', 'AU-4349'))
                st.session_state['logo_file'] = st.file_uploader("Company Logo", type=['png', 'jpg'])
                
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Logout", use_container_width=True):
                st.session_state.clear()
                st.rerun()

        # Page Routing
        if pg == "New Inspection": 
            inspection_page(ai)
        elif pg == "Manage Reports":
            manage_reports_page()
        elif pg == "Dashboard": 
            section_header("Inspection Dashboard", "dashboard")
            t_min, t_max = calculate_total_repairs(st.session_state['defects'])
            total_str = f"${t_min:,} - ${t_max:,}"

            c1,c2,c3 = st.columns(3)
            c1.metric("Defects Logged", len(st.session_state['defects']))
            c2.metric("Total Est. Cost", total_str)
            c3.metric("Property Type", st.session_state.get('property_type', 'N/A'))
            
            if st.session_state.get('maint_plan'):
                st.markdown("---")
                st.subheader("5-Year Maintenance Plan (AI)")
                st.markdown(st.session_state['maint_plan'])

        elif pg == "Finalize Report": 
            report_page(ai)
        elif pg == "Admin": 
            admin_page()

if __name__ == '__main__':
    main()
